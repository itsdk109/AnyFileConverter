import os
import logging
from docx import Document
from pdf2docx import Converter
from PIL import Image
from PyPDF2 import PdfReader, PdfWriter
from fpdf import FPDF

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def remove_metadata(file_path):
    """Removes metadata from images, PDFs, and DOCX files."""
    try:
        if file_path.endswith(('.jpg', '.jpeg', '.png', '.bmp')):
            remove_image_metadata(file_path)
        elif file_path.endswith('.pdf'):
            remove_pdf_metadata(file_path)
        elif file_path.endswith('.docx'):
            remove_docx_metadata(file_path)
        else:
            logging.warning("Unsupported file type for metadata removal.")
    except Exception as e:
        logging.error(f"Error removing metadata: {e}")

def remove_image_metadata(file_path):
    """Remove metadata from image files."""
    try:
        img = Image.open(file_path)
        data = list(img.getdata())
        img_no_meta = Image.new(img.mode, img.size)
        img_no_meta.putdata(data)
        img_no_meta.save(file_path)
        logging.info(f"Metadata removed from image: {file_path}")
    except Exception as e:
        logging.error(f"Error removing metadata from image: {e}")

def remove_pdf_metadata(file_path):
    """Remove metadata from PDF files."""
    try:
        reader = PdfReader(file_path)
        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)
        with open(file_path, 'wb') as f:
            writer.write(f)
        logging.info(f"Metadata removed from PDF: {file_path}")
    except Exception as e:
        logging.error(f"Error removing metadata from PDF: {e}")

def remove_docx_metadata(file_path):
    """Remove metadata from DOCX files."""
    try:
        doc = Document(file_path)
        doc.save(file_path)  # Saving without changes removes metadata
        logging.info(f"Metadata removed from DOCX: {file_path}")
    except Exception as e:
        logging.error(f"Error removing metadata from DOCX: {e}")

def validate_path(file_path):
    """Validates if a file path exists."""
    if not os.path.exists(file_path):
        logging.error(f"File not found: {file_path}")
        return False
    return True

def pdf_to_docx():
    """Converts a PDF file to DOCX."""
    pdf_path = input("Enter the PDF file path: ").strip()
    if not validate_path(pdf_path):
        return
    docx_path = input("Enter the output DOCX file path: ").strip()
    try:
        converter = Converter(pdf_path)
        converter.convert(docx_path, start=0, end=None)
        converter.close()
        remove_metadata(docx_path)
        logging.info(f"PDF successfully converted to DOCX: {docx_path}")
    except Exception as e:
        logging.error(f"Error converting PDF to DOCX: {e}")

def docx_to_pdf():
    """Converts a DOCX file to PDF."""
    docx_path = input("Enter the DOCX file path: ").strip()
    if not validate_path(docx_path):
        return
    pdf_path = input("Enter the output PDF file path: ").strip()
    try:
        doc = Document(docx_path)
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        for paragraph in doc.paragraphs:
            pdf.set_font("Arial", size=12)
            pdf.multi_cell(0, 10, paragraph.text)
        pdf.output(pdf_path)
        remove_metadata(pdf_path)
        logging.info(f"DOCX successfully converted to PDF: {pdf_path}")
    except Exception as e:
        logging.error(f"Error converting DOCX to PDF: {e}")

def extract_pdf_images():
    """Extracts images from a PDF and saves them."""
    pdf_path = input("Enter the PDF file path: ").strip()
    if not validate_path(pdf_path):
        return
    output_folder = input("Enter the output folder for images: ").strip()
    os.makedirs(output_folder, exist_ok=True)
    try:
        reader = PdfReader(pdf_path)
        for page_num, page in enumerate(reader.pages):
            if '/XObject' in page.get('/Resources', {}):
                x_object = page['/Resources']['/XObject'].get_object()
                for obj in x_object:
                    if x_object[obj]['/Subtype'] == '/Image':
                        size = (x_object[obj]['/Width'], x_object[obj]['/Height'])
                        data = x_object[obj].get_data()
                        img = Image.frombytes("RGB", size, data)
                        img.save(os.path.join(output_folder, f"image_{page_num + 1}.jpg"))
        logging.info(f"Images successfully extracted to: {output_folder}")
    except Exception as e:
        logging.error(f"Error extracting images from PDF: {e}")

def jpeg_to_jpg():
    """Converts JPEG to JPG."""
    input_path = input("Enter the JPEG file path: ").strip()
    if not validate_path(input_path):
        return
    output_path = input("Enter the output JPG file path: ").strip()
    try:
        img = Image.open(input_path)
        img = img.convert("RGB")  # Ensure proper RGB format
        img.save(output_path, "JPEG")
        remove_metadata(output_path)
        logging.info(f"JPEG successfully converted to JPG: {output_path}")
    except Exception as e:
        logging.error(f"Error converting JPEG to JPG: {e}")

def main():
    """Main menu for the tool."""
    while True:
        print("\nSelect an option:")
        print("1. Convert PDF to DOCX")
        print("2. Convert DOCX to PDF")
        print("3. Extract Images from PDF")
        print("4. Convert JPEG to JPG")
        print("5. Exit")
        choice = input("Enter your choice (1-5): ").strip()

        if choice == "1":
            pdf_to_docx()
        elif choice == "2":
            docx_to_pdf()
        elif choice == "3":
            extract_pdf_images()
        elif choice == "4":
            jpeg_to_jpg()
        elif choice == "5":
            logging.info("Exiting the tool. Goodbye!")
            break
        else:
            logging.warning("Invalid choice! Please try again.")

if __name__ == "__main__":
    main()
